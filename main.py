
import os
import time
import threading
import platform
import ctypes
import pyautogui
import tkinter as tk
from tkinter import messagebox, filedialog, StringVar
from tkinter import ttk
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageTk

# Modul 1: CursorHider - Verantwortlich für das Ausblenden des Mauszeigers
class CursorHider:
    def __init__(self):
        self.platform = platform.system()
        if self.platform == "Windows":
            self.user32 = ctypes.windll.user32
        elif self.platform == "Darwin":
            # macOS-spezifische Implementierung (falls erforderlich)
            pass
        elif self.platform == "Linux":
            # Linux-spezifische Implementierung (falls erforderlich)
            pass
        else:
            pass  # Andere Plattformen können hier hinzugefügt werden

    def __enter__(self):
        if self.platform == "Windows":
            # Cursor ausblenden
            self.user32.ShowCursor(False)
        # Weitere Plattform-spezifische Aktionen können hier hinzugefügt werden
        return self

    def __exit__(self, exc_type, exc_value, traceback):
        if self.platform == "Windows":
            # Cursor wieder einblenden
            self.user32.ShowCursor(True)
        # Weitere Plattform-spezifische Aktionen können hier hinzugefügt werden

# Modul 2: ScreenshotHandler - Verantwortlich für das Aufnehmen des Screenshot-Bereichs
class ScreenshotHandler:
    def __init__(self, app):
        self.app = app
        self.start_x = self.start_y = self.end_x = self.end_y = 0

    def select_area(self):
        # Hauptfenster ausblenden
        self.app.gui.root.withdraw()

        # Transparentes Vollbildfenster erstellen
        select_window = tk.Toplevel()
        select_window.attributes("-fullscreen", True)
        select_window.attributes("-alpha", 0.3)
        select_window.configure(background='grey')

        # Canvas zum Zeichnen des Rechtecks
        canvas = tk.Canvas(select_window, cursor="cross", bg="grey")
        canvas.pack(fill=tk.BOTH, expand=True)

        # Funktionen zum Verarbeiten der Mausereignisse
        def on_mouse_down(event):
            self.start_x, self.start_y = event.x, event.y
            canvas.delete("selection_rect")

        def on_mouse_move(event):
            canvas.delete("selection_rect")
            canvas.create_rectangle(self.start_x, self.start_y, event.x, event.y, outline='red', width=2, tag="selection_rect")

        def on_mouse_up(event):
            self.end_x, self.end_y = event.x, event.y
            select_window.destroy()
            self.app.gui.root.deiconify()
            self.app.gui.update_screenshot_fields()
            self.app.gui.show_preview()

        # Mausereignisse binden
        canvas.bind("<ButtonPress-1>", on_mouse_down)
        canvas.bind("<B1-Motion>", on_mouse_move)
        canvas.bind("<ButtonRelease-1>", on_mouse_up)

# Modul 3: DocumentHandler - Verantwortlich für die Erstellung und Verwaltung des Word-Dokuments
class DocumentHandler:
    def __init__(self):
        self.document = Document()

    def add_picture(self, picture_path, width=None, height=None):
        picture = self.document.add_picture(picture_path)
        if width:
            picture.width = width
        if height:
            picture.height = height

    def add_page_break(self):
        self.document.add_page_break()

    def save_document(self, save_path):
        self.document.save(save_path)

# Modul 4: GUI - Verantwortlich für die Benutzeroberfläche
class GUI:
    def __init__(self, root, app):
        self.root = root
        self.app = app

        # Stil für ttk
        style = ttk.Style()
        style.theme_use('clam')  # Moderne Optik

        # Hauptframe
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Abschnitt: Einstellungen
        settings_frame = ttk.LabelFrame(main_frame, text="Einstellungen", padding="15")
        settings_frame.grid(row=0, column=0, sticky="ew", padx=10, pady=10)

        # Anzahl der Screenshots
        ttk.Label(settings_frame, text="Anzahl der Screenshots:").grid(row=0, column=0, sticky="w", pady=5)
        self.entry_screenshot_count = ttk.Entry(settings_frame)
        self.entry_screenshot_count.insert(0, "5")  # Standardwert
        self.entry_screenshot_count.grid(row=0, column=1, pady=5, sticky="ew")

        # Verzögerung zwischen Screenshots
        ttk.Label(settings_frame, text="Verzögerung zwischen Screenshots (Sekunden):").grid(row=1, column=0, sticky="w", pady=5)
        self.entry_delay = ttk.Entry(settings_frame)
        self.entry_delay.insert(0, "2")  # Standardwert
        self.entry_delay.grid(row=1, column=1, pady=5, sticky="ew")

        # Scroll-Pixel pro Screenshot
        ttk.Label(settings_frame, text="Scroll-Pixel pro Screenshot:").grid(row=2, column=0, sticky="w", pady=5)
        self.entry_scroll_pixels = ttk.Entry(settings_frame)
        self.entry_scroll_pixels.insert(0, "365")  # Standardwert
        self.entry_scroll_pixels.grid(row=2, column=1, pady=5, sticky="ew")

        # Verzögerung vor Start
        ttk.Label(settings_frame, text="Verzögerung vor Start (Sekunden):").grid(row=3, column=0, sticky="w", pady=5)
        self.entry_start_delay = ttk.Entry(settings_frame)
        self.entry_start_delay.insert(0, "5")  # Standardwert
        self.entry_start_delay.grid(row=3, column=1, pady=5, sticky="ew")

        # Bildformat auswählen
        ttk.Label(settings_frame, text="Bildformat auswählen:").grid(row=4, column=0, sticky="w", pady=5)
        self.format_var = StringVar(value="PNG")
        self.format_menu = ttk.OptionMenu(settings_frame, self.format_var, "PNG", "PNG", "JPEG", "BMP")
        self.format_menu.grid(row=4, column=1, pady=5, sticky="ew")

        # Ausgabeverzeichnis
        ttk.Label(settings_frame, text="Ausgabeverzeichnis:").grid(row=5, column=0, sticky="w", pady=5)
        self.output_dir_var = StringVar(value="")
        self.output_dir_button = ttk.Button(settings_frame, text="Verzeichnis auswählen", command=self.app.select_output_directory)
        self.output_dir_button.grid(row=5, column=1, pady=5, sticky="ew")

        # Abschnitt: Screenshot-Bereich
        area_frame = ttk.LabelFrame(main_frame, text="Screenshot-Bereich", padding="15")
        area_frame.grid(row=1, column=0, sticky="ew", padx=10, pady=10)

        # X-Koordinate
        ttk.Label(area_frame, text="X-Koordinate des Screenshots:").grid(row=0, column=0, sticky="w", pady=5)
        self.entry_screenshot_x = ttk.Entry(area_frame)
        self.entry_screenshot_x.insert(0, "0")  # Standardwert
        self.entry_screenshot_x.grid(row=0, column=1, pady=5, sticky="ew")

        # Y-Koordinate
        ttk.Label(area_frame, text="Y-Koordinate des Screenshots:").grid(row=1, column=0, sticky="w", pady=5)
        self.entry_screenshot_y = ttk.Entry(area_frame)
        self.entry_screenshot_y.insert(0, "0")  # Standardwert
        self.entry_screenshot_y.grid(row=1, column=1, pady=5, sticky="ew")

        # Breite des Screenshots
        ttk.Label(area_frame, text="Breite des Screenshots:").grid(row=2, column=0, sticky="w", pady=5)
        self.entry_screenshot_width = ttk.Entry(area_frame)
        self.entry_screenshot_width.insert(0, "1024")  # Standardwert
        self.entry_screenshot_width.grid(row=2, column=1, pady=5, sticky="ew")

        # Höhe des Screenshots
        ttk.Label(area_frame, text="Höhe des Screenshots:").grid(row=3, column=0, sticky="w", pady=5)
        self.entry_screenshot_height = ttk.Entry(area_frame)
        self.entry_screenshot_height.insert(0, "768")  # Standardwert
        self.entry_screenshot_height.grid(row=3, column=1, pady=5, sticky="ew")

        # Bildgröße in Word-Dokument
        ttk.Label(area_frame, text="Breite des Bildes in Word (Inches):").grid(row=4, column=0, sticky="w", pady=5)
        self.entry_doc_image_width = ttk.Entry(area_frame)
        self.entry_doc_image_width.insert(0, "")  # Optional, kann leer bleiben
        self.entry_doc_image_width.grid(row=4, column=1, pady=5, sticky="ew")

        ttk.Label(area_frame, text="Höhe des Bildes in Word (Inches):").grid(row=5, column=0, sticky="w", pady=5)
        self.entry_doc_image_height = ttk.Entry(area_frame)
        self.entry_doc_image_height.insert(0, "")  # Optional, kann leer bleiben
        self.entry_doc_image_height.grid(row=5, column=1, pady=5, sticky="ew")

        # Button zur Auswahl des Screenshot-Bereichs
        self.select_area_button = ttk.Button(area_frame, text="Screenshot-Bereich auswählen", command=self.app.screenshot_handler.select_area)
        self.select_area_button.grid(row=6, column=0, columnspan=2, pady=10)

        # Abschnitt: Fortschritt
        progress_frame = ttk.LabelFrame(main_frame, text="Fortschritt", padding="15")
        progress_frame.grid(row=2, column=0, sticky="ew", padx=10, pady=10)

        progress_label = ttk.Label(progress_frame, text="Fortschritt:")
        progress_label.pack(anchor="w")
        self.progress = ttk.Progressbar(progress_frame, orient=tk.HORIZONTAL, length=500, mode='determinate')
        self.progress.pack(pady=5, fill="x")

        # Abschnitt: Buttons
        buttons_frame = ttk.Frame(main_frame, padding="15")
        buttons_frame.grid(row=3, column=0, sticky="ew", padx=10, pady=10)

        self.start_button = ttk.Button(buttons_frame, text="Start", command=self.app.start_capture_thread)
        self.start_button.pack(side="left", expand=True, fill="x", padx=5)

        self.stop_button = ttk.Button(buttons_frame, text="Stop", command=self.app.stop_capture)
        self.stop_button.pack(side="left", expand=True, fill="x", padx=5)

        # Anpassen der Spaltengewichte für gleichmäßige Verteilung
        settings_frame.columnconfigure(0, weight=1)
        settings_frame.columnconfigure(1, weight=2)
        area_frame.columnconfigure(0, weight=1)
        area_frame.columnconfigure(1, weight=2)
        progress_frame.columnconfigure(0, weight=1)
        buttons_frame.columnconfigure(0, weight=1)
        buttons_frame.columnconfigure(1, weight=1)

    def update_screenshot_fields(self):
        # Bereich berechnen
        x1 = min(self.app.screenshot_handler.start_x, self.app.screenshot_handler.end_x)
        y1 = min(self.app.screenshot_handler.start_y, self.app.screenshot_handler.end_y)
        width = abs(self.app.screenshot_handler.end_x - self.app.screenshot_handler.start_x)
        height = abs(self.app.screenshot_handler.end_y - self.app.screenshot_handler.start_y)

        # Eingabefelder aktualisieren
        self.entry_screenshot_x.delete(0, tk.END)
        self.entry_screenshot_x.insert(0, str(x1))
        self.entry_screenshot_y.delete(0, tk.END)
        self.entry_screenshot_y.insert(0, str(y1))
        self.entry_screenshot_width.delete(0, tk.END)
        self.entry_screenshot_width.insert(0, str(width))
        self.entry_screenshot_height.delete(0, tk.END)
        self.entry_screenshot_height.insert(0, str(height))

    def show_preview(self):
        try:
            x1 = int(self.entry_screenshot_x.get())
            y1 = int(self.entry_screenshot_y.get())
            width = int(self.entry_screenshot_width.get())
            height = int(self.entry_screenshot_height.get())
            preview_screenshot = pyautogui.screenshot(region=(x1, y1, width, height))
            preview_window = tk.Toplevel(self.root)
            preview_window.title("Vorschau des Screenshot-Bereichs")
            preview_image = ImageTk.PhotoImage(preview_screenshot)
            label = ttk.Label(preview_window, image=preview_image)
            label.image = preview_image  # Referenz speichern
            label.pack(padx=10, pady=10)
        except Exception as e:
            messagebox.showerror("Fehler", f"Vorschau konnte nicht erstellt werden:\n{e}")

# Modul 5: SnapToWordApp - Koordiniert die Interaktion zwischen GUI, Screenshot- und Dokumenten-Handler
class SnapToWordApp:
    def __init__(self, root):
        self.root = root
        self.document_handler = DocumentHandler()
        self.screenshot_handler = ScreenshotHandler(self)
        self.gui = GUI(root, self)
        self.stop_flag = False

    def start_capture_thread(self):
        threading.Thread(target=self.capture_screenshots, daemon=True).start()

    def capture_screenshots(self):
        try:
            self.stop_flag = False  # Stop-Flag zurücksetzen

            # Einstellungen einlesen
            screenshot_count = int(self.gui.entry_screenshot_count.get())
            delay = float(self.gui.entry_delay.get())
            scroll_pixels = int(self.gui.entry_scroll_pixels.get())
            image_format = self.gui.format_var.get().lower()
            output_directory = self.gui.output_dir_var.get()

            # Bildgröße in Word einlesen
            doc_image_width = self.gui.entry_doc_image_width.get()
            doc_image_height = self.gui.entry_doc_image_height.get()

            if doc_image_width:
                doc_image_width = Inches(float(doc_image_width))  # Umwandeln in Inches
            else:
                doc_image_width = None

            if doc_image_height:
                doc_image_height = Inches(float(doc_image_height))  # Umwandeln in Inches
            else:
                doc_image_height = None

            if not output_directory:
                messagebox.showerror("Fehler", "Bitte ein Ausgabeverzeichnis auswählen.")
                return

            if not os.path.exists(output_directory):
                os.makedirs(output_directory)

            # Screenshot-Bereich einlesen
            left = int(self.gui.entry_screenshot_x.get())
            top = int(self.gui.entry_screenshot_y.get())
            width = int(self.gui.entry_screenshot_width.get())
            height = int(self.gui.entry_screenshot_height.get())

            # Fortschrittsleiste initialisieren
            self.gui.progress['value'] = 0
            self.gui.progress['maximum'] = screenshot_count

            # Verzögerung vor Start
            start_delay = float(self.gui.entry_start_delay.get())
            time.sleep(start_delay)

            # Screenshots aufnehmen mit Cursor ausblenden
            with CursorHider():
                for i in range(screenshot_count):
                    if self.stop_flag:
                        break

                    # Screenshot des angegebenen Bereichs aufnehmen
                    screenshot = pyautogui.screenshot(region=(left, top, width, height))
                    screenshot_file = os.path.join(output_directory, f"screenshot_{i+1}.{image_format}")
                    screenshot.save(screenshot_file, format=image_format.upper())
                    print(f"Screenshot {i+1} gespeichert als {screenshot_file}")

                    # Screenshot in Word-Dokument einfügen
                    if doc_image_width or doc_image_height:
                        self.document_handler.add_picture(screenshot_file, width=doc_image_width, height=doc_image_height)
                    else:
                        self.document_handler.add_picture(screenshot_file)

                    if i < screenshot_count - 1:
                        self.document_handler.add_page_break()

                    # Temporäre Datei löschen
                    os.remove(screenshot_file)

                    # Mauszeiger in die Mitte des Screenshot-Bereichs bewegen
                    pyautogui.moveTo(left + width / 2, top + height / 2)
                    # Scrollen
                    pyautogui.scroll(-scroll_pixels)

                    # Fortschrittsleiste aktualisieren
                    self.gui.progress['value'] = i + 1
                    self.root.update_idletasks()

                    # Wartezeit zwischen den Screenshots
                    time.sleep(delay)

            # Word-Datei speichern
            if not self.stop_flag:
                save_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                         filetypes=[("Word Datei", "*.docx")],
                                                         initialdir=output_directory)
                if save_path:
                    self.document_handler.save_document(save_path)
                    messagebox.showinfo("Erfolg", f"Word-Datei gespeichert: {save_path}")

        except Exception as e:
            messagebox.showerror("Fehler", str(e))

    def stop_capture(self):
        self.stop_flag = True
        messagebox.showinfo("Gestoppt", "Screenshot-Aufnahme gestoppt.")

    def select_output_directory(self):
        directory = filedialog.askdirectory()
        if directory:
            self.gui.output_dir_var.set(directory)

# Hauptfunktion zum Starten der Anwendung
def main():
    root = tk.Tk()
    app = SnapToWordApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
